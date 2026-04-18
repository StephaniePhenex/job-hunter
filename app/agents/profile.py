"""User profile for relevance scoring — loaded from user_profile.yaml."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

from app.core.config import get_settings

logger = logging.getLogger(__name__)

_REPO_ROOT = Path(__file__).resolve().parents[2]
_DEFAULT_YAML_PATH = _REPO_ROOT / "user_profile.yaml"


def _normalize_resume_id(stem: str) -> str:
    """Derive resume id from file stem: lowercase, [a-z0-9_-], max 64 chars."""
    s = stem.strip().lower()
    s = re.sub(r"[^a-z0-9_-]+", "-", s)
    s = re.sub(r"-+", "-", s).strip("-")
    if not s:
        return ""
    return s[:64]


def _read_resume_file_text(path: Path) -> str:
    """Read plain text from .md / .txt / .docx (Word 2007+)."""
    suf = path.suffix.lower()
    if suf in (".md", ".txt"):
        return path.read_text(encoding="utf-8").strip()
    if suf == ".docx":
        try:
            from docx import Document
        except ImportError:
            logger.warning(
                "Skipping %s: python-docx not installed (pip install python-docx)",
                path.name,
            )
            return ""
        try:
            doc = Document(str(path))
            parts: list[str] = []
            for para in doc.paragraphs:
                t = (para.text or "").strip()
                if t:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = " | ".join(
                        (cell.text or "").strip() for cell in row.cells if (cell.text or "").strip()
                    )
                    if cells:
                        parts.append(cells)
            return "\n".join(parts).strip()
        except Exception:
            logger.exception("Could not read Word resume %s", path)
            return ""
    return ""


def load_resume_entries_from_disk(resume_dir: Path) -> list[ResumeEntry]:
    """Load *.md / *.txt / *.docx from resume_dir; id from filename stem (normalized). Later file wins on id clash."""
    if not resume_dir.is_dir():
        return []
    paths = sorted(
        list(resume_dir.glob("*.md"))
        + list(resume_dir.glob("*.txt"))
        + list(resume_dir.glob("*.docx")),
        key=lambda p: p.name.lower(),
    )
    by_id: dict[str, ResumeEntry] = {}
    for path in paths:
        stem = path.stem
        if stem.lower() == "readme":
            continue
        if stem.startswith("."):
            continue
        try:
            content = _read_resume_file_text(path)
        except OSError:
            logger.warning("Could not read resume file %s", path)
            continue
        if not content:
            continue
        rid = _normalize_resume_id(stem)
        if not rid:
            continue
        if rid in by_id:
            logger.warning(
                "Duplicate resume id %s from file %s (overwrites previous file)",
                rid,
                path.name,
            )
        by_id[rid] = ResumeEntry(id=rid, content=content)
    return sorted(by_id.values(), key=lambda r: r.id)


def _merge_disk_resumes(
    profile: UserProfile,
    *,
    project_root: Path,
    resume_dir: Path,
    enabled: bool,
) -> None:
    if not enabled:
        return
    disk = load_resume_entries_from_disk(resume_dir)
    if not disk:
        return
    yaml_first = list(profile.resumes)
    yaml_ids = [r.id for r in yaml_first]
    by_id: dict[str, ResumeEntry] = {r.id: r for r in yaml_first}
    for r in disk:
        by_id[r.id] = r
    ordered: list[ResumeEntry] = []
    for rid in yaml_ids:
        if rid in by_id:
            ordered.append(by_id[rid])
    disk_only = sorted(
        [r for r in disk if r.id not in yaml_ids],
        key=lambda x: x.id,
    )
    ordered.extend(disk_only)
    profile.resumes = ordered


class ResumeEntry(BaseModel):
    """One resume in the library. `id` is the routing key; `content` is plain text."""

    id: str = Field(min_length=1, max_length=64)
    content: str = Field(min_length=1)


class UserProfile(BaseModel):
    """Target intern seeker preferences passed to the LLM for scoring."""

    location_focus: str = Field(default="Canada", description="Geographic focus")
    term: str = Field(default="Fall internship", description="Preferred internship term")
    interests: list[str] = Field(
        default_factory=lambda: [
            "Web3",
            "Software engineering",
            "Media / content",
            "Tech + culture intersection",
        ]
    )
    resumes: list[ResumeEntry] = Field(default_factory=list)
    # Optimize agent anti-hallucination guards
    verified_skills: list[str] = Field(default_factory=list)
    never_claim: list[str] = Field(default_factory=list)


def _load_yaml(path: Path) -> dict[str, Any]:
    try:
        import yaml  # PyYAML (already in requirements via pydantic-settings extras or explicit)
    except ImportError:
        logger.warning("PyYAML not installed; using default profile. Run: pip install pyyaml")
        return {}
    try:
        with open(path, encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except FileNotFoundError:
        logger.debug("user_profile.yaml not found at %s; using defaults", path)
        return {}
    except Exception:
        logger.exception("Failed to parse %s; using defaults", path)
        return {}


def load_user_profile(
    path: Path | None = None,
    *,
    project_root: Path | None = None,
    resume_files_dir: str | None = None,
    resume_files_enabled: bool | None = None,
) -> UserProfile:
    """Load UserProfile (profile + resumes sections) from YAML, then merge `data/resumes` files.

    Disk files (*.docx / *.md / *.txt) override YAML entries with the same `id`. Restart uvicorn to reload.
    """
    resolved = path or _DEFAULT_YAML_PATH
    root = project_root or _REPO_ROOT
    data = _load_yaml(resolved)
    raw: dict[str, Any] = dict(data.get("profile") or {})
    raw_resumes = data.get("resumes")
    raw["resumes"] = raw_resumes if isinstance(raw_resumes, list) else []
    for top_key in ("verified_skills", "never_claim"):
        val = data.get(top_key)
        if isinstance(val, list):
            raw[top_key] = val
    profile = UserProfile.model_validate(raw)

    settings = get_settings()
    enabled = settings.resume_files_enabled if resume_files_enabled is None else resume_files_enabled
    subdir = settings.resume_files_dir if resume_files_dir is None else resume_files_dir
    resume_dir = (root / subdir).resolve()
    _merge_disk_resumes(profile, project_root=root, resume_dir=resume_dir, enabled=enabled)
    return profile


def resumes_as_dicts(profile: UserProfile) -> list[dict]:
    """Return resumes in the format expected by the analyze prompt: [{id, content}]."""
    return [{"id": r.id, "content": r.content} for r in profile.resumes]


DEFAULT_USER_PROFILE = load_user_profile()
